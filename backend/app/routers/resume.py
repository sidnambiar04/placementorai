"""
Resume analysis and optimization router.
"""
import io
import json
import re
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import Optional
from app.ai.gemini_client import client, MODEL, SAFETY_SETTINGS, _extract_json
from app.ai.prompts.resume_prompt import build_resume_analysis_prompt, build_resume_optimize_prompt

router = APIRouter()


def _extract_text_from_file(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Extract text from PDF, DOCX, or image file."""
    fname = filename.lower()
    
    if content_type == "application/pdf" or fname.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text.strip()
        except Exception as e:
            print(f"PDF parse error: {e}")
            return ""
    
    elif fname.endswith((".docx",)) or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import zipfile
            from xml.etree import ElementTree
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                with z.open("word/document.xml") as doc_xml:
                    tree = ElementTree.parse(doc_xml)
                    root = tree.getroot()
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    texts = [node.text for node in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if node.text]
                    return " ".join(texts).strip()
        except Exception as e:
            print(f"DOCX parse error: {e}")
            return ""
    
    elif fname.endswith((".doc",)):
        # Basic .doc support - return placeholder
        return "DOC format detected - basic parsing applied"
    
    elif content_type.startswith("image/") or fname.endswith((".png", ".jpg", ".jpeg")):
        # For images, use Gemini Vision
        return "[IMAGE_RESUME]"
    
    return ""


async def _analyze_with_gemini(file_bytes: bytes, filename: str, content_type: str,
                                target_role: str, job_description: str) -> dict:
    """Run resume analysis through Gemini."""
    fname = filename.lower()
    
    # For images, use multi-modal Gemini
    if content_type.startswith("image/") or fname.endswith((".png", ".jpg", ".jpeg")):
        from google.genai import types as genai_types
        
        mime_map = {"image/png": "image/png", "image/jpeg": "image/jpeg", "image/jpg": "image/jpeg"}
        img_mime = mime_map.get(content_type, "image/png")
        
        jd_section = f"\n\nJob Description: {job_description}" if job_description.strip() else ""
        prompt_text = build_resume_analysis_prompt("[See image above]", target_role, job_description)
        
        response = client.models.generate_content(
            model=MODEL,
            contents=[
                genai_types.Part.from_bytes(data=file_bytes, mime_type=img_mime),
                genai_types.Part.from_text(text=prompt_text),
            ],
            config={
                "response_mime_type": "application/json",
                "safety_settings": SAFETY_SETTINGS,
            }
        )
    else:
        resume_text = _extract_text_from_file(file_bytes, filename, content_type)
        if not resume_text.strip():
            raise ValueError("Could not extract text from the uploaded file. Please upload a PDF, DOCX, image, or text file.")
        
        prompt = build_resume_analysis_prompt(resume_text, target_role, job_description)
        response = client.models.generate_content(
            model=MODEL,
            contents=prompt,
            config={
                "response_mime_type": "application/json",
                "safety_settings": SAFETY_SETTINGS,
            }
        )

    raw = response.text
    print(f"DEBUG RAW ANALYSIS (first 400): {raw[:400]}")
    
    cleaned = _extract_json(raw)
    result = json.loads(cleaned)
    
    # Ensure all metrics keys exist for the Spider Chart to prevent chart break
    if "metrics" not in result or not isinstance(result["metrics"], dict):
        result["metrics"] = {}
    
    for m in ["keywords", "format", "completeness", "role", "impact"]:
        if m not in result["metrics"]:
            result["metrics"][m] = 0
            
    return result


@router.post("/analyze")
async def analyze_resume(
    file: UploadFile = File(...),
    target_role: str = Form(default="Software Engineer"),
    job_description: str = Form(default=""),
):
    try:
        file_bytes = await file.read()
        analysis = await _analyze_with_gemini(
            file_bytes=file_bytes,
            filename=file.filename or "resume",
            content_type=file.content_type or "",
            target_role=target_role,
            job_description=job_description,
        )
        return {"analysis": analysis}

    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        err_str = str(e).upper()
        print(f"ERROR analyze_resume: {e}")
        if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str or "QUOTA" in err_str:
            print("FALLBACK: Using mock data for resume analysis due to quota limit.")
            return {"analysis": get_mock_resume_analysis(target_role), "is_mock": True}
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


def get_mock_resume_analysis(role: str):
    return {
        "atsScore": 75,
        "potential": "Strong",
        "confidence": "HIGH",
        "metrics": {
            "keywords": 85,
            "format": 90,
            "completeness": 80,
            "role": 70,
            "impact": 75
        },
        "atsKillers": [
            "Missing specific tech stack keywords for this role",
            "Slightly wordy summaries in experience sections"
        ],
        "recruiterReality": "You have the core skills, but your resume doesn't shout 'Expert' yet. Focus on quantifiable achievements.",
        "keywordGap": {
            "important": ["Docker", "Kubernetes", "Redis"],
            "nice_to_have": ["GraphQL", "Unit Testing"]
        },
        "optimizedBullets": [
            {
                "label": "Clarity & Impact",
                "original": "Worked on the frontend using React.",
                "optimized": "Engineered high-performance React components, reducing load time by 40%.",
                "severity": "HIGH",
                "improvements": ["Action-oriented", "Quantified results"]
            }
        ],
        "finalVerdict": {
            "status": "Ready for Review",
            "fixTime": "2 hours",
            "points": [
                {"type": "green", "text": "Clean, ATS-friendly layout"},
                {"type": "yellow", "text": "Needs more power verbs"},
                {"type": "red", "text": "Missing cloud infrastructure mentions"}
            ]
        }
    }


@router.post("/optimize")
async def optimize_resume(
    file: UploadFile = File(...),
    target_role: str = Form(default="Software Engineer"),
    analysis: str = Form(default="{}"),
):
    """Generate an optimized DOCX resume based on analysis feedback."""
    try:
        file_bytes = await file.read()
        fname = (file.filename or "resume").lower()
        content_type = file.content_type or ""

        # Extract original text
        if content_type.startswith("image/") or fname.endswith((".png", ".jpg", ".jpeg")):
            resume_text = "Please improve this resume for the target role."
        else:
            resume_text = _extract_text_from_file(file_bytes, fname, content_type)
            if not resume_text.strip():
                resume_text = "Resume content could not be extracted. Create an optimized template."

        analysis_dict = json.loads(analysis)
        prompt = build_resume_optimize_prompt(resume_text, target_role, analysis_dict)

        response = client.models.generate_content(
            model=MODEL,
            contents=prompt,
            config={"safety_settings": SAFETY_SETTINGS}
        )

        optimized_text = response.text

        # Build a DOCX from the optimized text
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(f"Optimized Resume — {target_role}", 0)
            for line in optimized_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("- ") or line.startswith("• "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="Optimized_Resume.docx"'}
            )
        except ImportError:
            # If python-docx is not installed, return plain text
            text_buf = io.BytesIO(optimized_text.encode("utf-8"))
            return StreamingResponse(
                text_buf,
                media_type="text/plain",
                headers={"Content-Disposition": 'attachment; filename="Optimized_Resume.txt"'}
            )

    except Exception as e:
        print(f"ERROR optimize_resume: {e}")
        raise HTTPException(status_code=500, detail=f"Optimization failed: {str(e)}")
